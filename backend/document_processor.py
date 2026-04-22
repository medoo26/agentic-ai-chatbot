import os
import re
import inspect
import asyncio
import time
from pathlib import Path
from typing import Dict, List, Tuple, Any
from html import unescape

from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import HTMLHeaderTextSplitter


class DocumentProcessor:
    def __init__(self, rag_system, llm_service=None):
        self.rag_system = rag_system
        self.llm_service = llm_service

        # OCR optional (default OFF)
        self.enable_ocr = (os.getenv("ENABLE_OCR", "0").strip() == "1")
        self.tesseract_cmd = (os.getenv("TESSERACT_CMD") or "").strip()

        # Folder for converted HTML (stored as .txt)
        processed_dir = (os.getenv("PROCESSED_HTML_DIR") or "processed_html").strip()
        self.processed_html_dir = Path(processed_dir)
        self.processed_html_dir.mkdir(parents=True, exist_ok=True)

        # Chunking settings (for indexing only)
        self.chunk_size = int(os.getenv("HTML_CHUNK_SIZE", "1800"))
        self.chunk_overlap = int(os.getenv("HTML_CHUNK_OVERLAP", "350"))

        # LLM conversion protection
        self.html_llm_timeout = int(os.getenv("HTML_LLM_TIMEOUT", "60"))

        # final hard cap
        self.html_hard_max_chars = int(os.getenv("HTML_HARD_MAX_CHARS", "250000"))

        # Segment large docs for LLM conversion
        self.html_segment_chars = int(os.getenv("HTML_SEGMENT_CHARS", "12000"))
        self.html_segment_overlap = int(os.getenv("HTML_SEGMENT_OVERLAP", "400"))

        # Merge tiny chunks so RAG doesn't retrieve headings only
        self.min_chunk_chars = int(os.getenv("HTML_MIN_CHUNK_CHARS", "250"))

        # Legal HTML: split chunks at h1/h2 only (full مادة including h3/p). LEGAL_SPLIT_ON_H3=1 restores old h3 boundaries.
        self.legal_split_on_h3 = (os.getenv("LEGAL_SPLIT_ON_H3") or "0").strip() == "1"
        # If one h2 section exceeds this (plain-text chars), split into parts with same headers (0 = disabled).
        self.legal_h2_max_chars = int((os.getenv("LEGAL_H2_MAX_CHARS") or "0").strip() or "0")

        # PDF_HTML_MODE=gemini_file: PDF→HTML via Gemini Files API + to_structured_html_from_pdf (no pypdf for HTML).
        # Requires GEMINI_API_KEY or GOOGLE_API_KEY; optional GEMINI_MODEL_NAME, HTML_MAX_TOKENS.

    def _trace(self, request_id: str, step: str, t0: float, extra: str = "") -> None:
        elapsed = time.perf_counter() - t0
        suffix = f" | {extra}" if extra else ""
        print(f"[TRACE][DOC_PROCESS][{request_id}] {step} | +{elapsed:.2f}s{suffix}")

    # =========================================================
    # Helpers: clean stored filename -> original filename
    # =========================================================
    def _clean_original_name(self, stored_basename: str) -> str:
        name = (stored_basename or "").strip()
        if not name:
            return ""

        m = re.match(
            r"^[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}_(.+)$",
            name,
        )
        if m:
            return (m.group(1) or "").strip()

        parts = name.split("_", 1)
        if len(parts) == 2 and len(parts[0]) >= 20:
            return parts[1].strip()

        return name

    async def _maybe_await(self, value: Any) -> Any:
        if inspect.isawaitable(value):
            return await value
        return value

    # =========================================================
    # Save converted HTML (as .txt)
    # =========================================================
    def _save_converted_html_as_txt(self, original_filename: str, html: str) -> Tuple[str, str]:
        stem = Path(original_filename).stem
        out_name = f"{stem}__html.txt"
        out_path = self.processed_html_dir / out_name
        out_path.write_text(html or "", encoding="utf-8")
        return str(out_path), out_name

    # =========================================================
    # Split long raw text into segments for LLM conversion
    # =========================================================
    def _split_text_segments(self, text: str) -> List[str]:
        t = (text or "").strip()
        if not t:
            return []

        seg = max(1500, self.html_segment_chars)
        ov = max(0, min(self.html_segment_overlap, seg - 1))
        step = max(1, seg - ov)

        if len(t) <= seg:
            return [t]

        out: List[str] = []
        for i in range(0, len(t), step):
            part = t[i : i + seg].strip()
            if part:
                out.append(part)
            if i + seg >= len(t):
                break
        return out

    # =========================================================
    # Robust HTML body extraction (prevents multiple doctypes)
    # =========================================================
    def _strip_outer_html(self, html: str) -> str:
        h = (html or "").strip()
        if not h:
            return ""

        h = re.sub(r"(?is)<!doctype[^>]*>", "", h).strip()
        h = re.sub(r"(?is)<html\b[^>]*>", "", h).strip()
        h = re.sub(r"(?is)</html\s*>", "", h).strip()
        h = re.sub(r"(?is)<head\b[^>]*>.*?</head\s*>", "", h).strip()

        m = re.search(r"(?is)<body\b[^>]*>(.*?)</body\s*>", h)
        if m:
            return (m.group(1) or "").strip()

        h = re.sub(r"(?is)<body\b[^>]*>", "", h).strip()
        h = re.sub(r"(?is)</body\s*>", "", h).strip()
        return h.strip()

    def _wrap_html_document(self, body_inner: str, title: str) -> str:
        t = (title or "Document").strip() or "Document"
        b = (body_inner or "").strip()
        return f"""<!doctype html>
<html lang="ar">
<head>
  <meta charset="utf-8">
  <title>{t}</title>
</head>
<body>
{b}
</body>
</html>
""".strip()

    # =========================================================
    # HTML -> plain text (for embeddings)
    # =========================================================
    def _html_to_plain(self, html: str) -> str:
        x = (html or "").strip()
        if not x:
            return ""
        x = re.sub(r"(?is)<br\s*/?>", "\n", x)
        x = re.sub(r"(?is)</p\s*>", "\n\n", x)
        x = re.sub(r"(?is)</h[1-6]\s*>", "\n", x)
        x = re.sub(r"(?is)</td\s*>", " | ", x)
        x = re.sub(r"(?is)</th\s*>", " | ", x)
        x = re.sub(r"(?is)</tr\s*>", "\n", x)
        x = re.sub(r"(?is)<[^>]+>", " ", x)
        x = re.sub(r"[ \t]+", " ", x)
        x = re.sub(r"\n{3,}", "\n\n", x)
        return x.strip()

    # =========================================================
    # Header / Footer removal helpers
    # =========================================================
    def _is_page_number_line(self, line: str) -> bool:
        s = (line or "").strip()
        if not s:
            return False

        s = s.replace("الصفحة", "").replace("صفحة", "").strip()
        s = s.translate(str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789"))

        return bool(re.fullmatch(r"\d{1,4}", s))

    def _normalize_header_footer_line(self, line: str) -> str:
        s = (line or "").strip()
        if not s:
            return ""

        s = s.translate(str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789"))
        s = re.sub(r"\d+", "#", s)
        s = re.sub(r"\s+", " ", s).strip()
        return s

    def _collect_repeated_edge_lines(
        self,
        pdf_pages: List[Dict[str, Any]],
        edge_line_count: int = 3,
        min_repeat: int = 2,
    ) -> Tuple[set, set]:
        """
        يجمع السطور المتكررة في أعلى وأسفل الصفحات لاستخدامها كهيدر/فوتر.
        """
        from collections import Counter

        top_counter = Counter()
        bottom_counter = Counter()

        for page in pdf_pages or []:
            text = (page.get("text") or "").strip()
            if not text:
                continue

            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            if not lines:
                continue

            top_lines = lines[:edge_line_count]
            bottom_lines = lines[-edge_line_count:]

            for ln in top_lines:
                norm = self._normalize_header_footer_line(ln)
                if norm and len(norm) >= 2:
                    top_counter[norm] += 1

            for ln in bottom_lines:
                norm = self._normalize_header_footer_line(ln)
                if norm and len(norm) >= 2:
                    bottom_counter[norm] += 1

        repeated_top = {k for k, v in top_counter.items() if v >= min_repeat}
        repeated_bottom = {k for k, v in bottom_counter.items() if v >= min_repeat}

        return repeated_top, repeated_bottom

    def _remove_header_footer_from_page_text(
        self,
        text: str,
        repeated_top: set,
        repeated_bottom: set,
        edge_line_count: int = 3,
    ) -> str:
        if not text:
            return ""

        lines = [ln.rstrip() for ln in text.splitlines()]
        if not lines:
            return ""

        cleaned = []

        for idx, line in enumerate(lines):
            raw = (line or "").strip()
            if not raw:
                continue

            norm = self._normalize_header_footer_line(raw)

            # إزالة أرقام الصفحات المنفصلة
            if self._is_page_number_line(raw):
                continue

            # إزالة الهيدر المتكرر من أعلى الصفحة فقط
            if idx < edge_line_count and norm in repeated_top:
                continue

            # إزالة الفوتر المتكرر من أسفل الصفحة فقط
            if idx >= max(0, len(lines) - edge_line_count) and norm in repeated_bottom:
                continue

            cleaned.append(raw)

        out = "\n".join(cleaned)
        out = self._clean_text_keep_tables(out)
        return out

    # =========================================================
    # Table normalization + strict legal split (H1/H2/H3)
    # =========================================================
    _RE_TABLE = re.compile(r"(?is)<table\b[^>]*>.*?</table>")
    _RE_TR = re.compile(r"(?is)<tr\b[^>]*>(.*?)</tr>")
    _RE_CELL = re.compile(r"(?is)<t[dh]\b[^>]*>(.*?)</t[dh]>")
    _RE_TAG = re.compile(r"(?is)<[^>]+>")
    _RE_PAGE_ATTR = re.compile(r'(?is)data-page\s*=\s*"(\d+)"')

    def _html_fragment_to_text(self, html_fragment: str) -> str:
        t = (html_fragment or "").strip()
        if not t:
            return ""
        t = re.sub(r"(?is)<br\s*/?>", "\n", t)
        t = re.sub(r"(?is)</p\s*>", "\n\n", t)
        t = re.sub(r"(?is)</h[1-6]\s*>", "\n", t)
        t = re.sub(r"(?is)</li\s*>", "\n", t)
        t = self._RE_TAG.sub(" ", t)
        t = unescape(t)
        t = re.sub(r"[ \t]+", " ", t)
        t = re.sub(r"\n{3,}", "\n\n", t)
        return t.strip()

    def _normalize_table_to_text(self, table_html: str, table_id: str) -> Tuple[str, str]:
        rows: List[List[str]] = []
        for tr in self._RE_TR.findall(table_html or ""):
            cells = [self._html_fragment_to_text(c) for c in self._RE_CELL.findall(tr or "")]
            cells = [c for c in cells if c]
            if cells:
                rows.append(cells)

        if not rows:
            return "", ""

        header = rows[0]
        body = rows[1:] if len(rows) > 1 else []
        schema = " | ".join(header)
        row_texts: List[str] = []
        for r in body:
            cols: List[str] = []
            max_len = min(len(header), len(r))
            for i in range(max_len):
                cols.append(f"{header[i]}={r[i]}")
            if not cols and r:
                cols = [f"col_{i+1}={val}" for i, val in enumerate(r)]
            row_texts.append("[" + ", ".join(cols) + "]")

        canonical = f"Table: {table_id}; schema: {schema}"
        if row_texts:
            canonical += " | " + " | ".join(row_texts)
        return canonical.strip(), schema.strip()

    def _normalize_html_tables(self, html_doc: str) -> Tuple[str, List[Dict[str, Any]]]:
        html_doc = (html_doc or "").strip()
        if not html_doc:
            return "", []

        table_meta: List[Dict[str, Any]] = []

        def _replace_table(match):
            idx = len(table_meta) + 1
            table_id = f"tbl_{idx}"
            table_html = (match.group(0) or "").strip()
            table_text, schema = self._normalize_table_to_text(table_html, table_id)
            if not table_text:
                return ""
            table_meta.append(
                {
                    "table_id": table_id,
                    "table_text": table_text,
                    "table_schema": schema,
                }
            )
            return f"<p>{table_text}</p>"

        normalized = self._RE_TABLE.sub(_replace_table, html_doc)
        return normalized, table_meta

    def _section_type_from_headers(self, headers: Dict[str, str], body_text: str) -> str:
        article = (headers.get("article") or "").strip()
        if self._is_table_section(body_text):
            return "table"
        txt = (body_text or "").strip().lower()
        if article:
            return "article"
        if (headers.get("level_2") or "").strip():
            return "h2_section"
        if (headers.get("level_1") or "").strip():
            return "chapter"
        return "body"

    def _looks_like_table_text(self, text: str) -> bool:
        t = (text or "").strip()
        if not t:
            return False
        tl = t.lower()
        if "table:" in tl:
            return True
        # Markdown table signature: has pipes and a separator row.
        if "|" in t and re.search(r"^\s*\|?\s*[-:]{3,}", t, flags=re.MULTILINE):
            return True
        return False

    def _is_table_section(self, text: str) -> bool:
        """
        Conservative detector:
        mark as table only when the section is table-only/predominantly table.
        This avoids classifying full h2/article sections as table just because they
        contain one embedded table line.
        """
        t = (text or "").strip()
        if not t:
            return False

        # Canonical table blocks generated by _normalize_html_tables are single-line-like.
        if t.lower().startswith("table:"):
            return True

        lines = [ln.strip() for ln in t.splitlines() if ln.strip()]
        if not lines:
            return False

        tableish = 0
        non_tableish = 0
        for ln in lines:
            lnl = ln.lower()
            if "table:" in lnl:
                tableish += 1
                continue
            if "|" in ln and re.search(r"^\s*\|?\s*[-:]{3,}", ln):
                tableish += 1
                continue
            non_tableish += 1

        # Consider it a table section only when almost everything is table-like.
        return tableish > 0 and non_tableish == 0

    _RE_FIRST_H3 = re.compile(r"(?is)<h3\b[^>]*>(.*?)</h3\s*>")

    def _derive_first_h3_plain(self, section_html: str) -> str:
        m = self._RE_FIRST_H3.search(section_html or "")
        if not m:
            return ""
        return self._html_to_plain(m.group(1)).strip()

    def _split_plain_into_parts(self, text: str, max_chars: int) -> List[str]:
        """Split plain text at paragraph boundaries; each part at most max_chars (fallback hard cut)."""
        t = (text or "").strip()
        if not t or max_chars <= 0:
            return [t] if t else []
        paras = [p.strip() for p in re.split(r"\n\s*\n+", t) if p.strip()]
        if not paras:
            return [t[:max_chars]]
        parts: List[str] = []
        buf: List[str] = []
        cur_len = 0
        for p in paras:
            add_len = len(p) if not buf else len(p) + 2
            if buf and cur_len + add_len > max_chars:
                parts.append("\n\n".join(buf))
                buf = [p]
                cur_len = len(p)
                continue
            buf.append(p)
            cur_len += add_len
        if buf:
            parts.append("\n\n".join(buf))
        # Handle single huge paragraph
        out: List[str] = []
        for chunk in parts:
            if len(chunk) <= max_chars:
                out.append(chunk)
                continue
            i = 0
            while i < len(chunk):
                out.append(chunk[i : i + max_chars])
                i += max_chars
        return [x for x in out if x.strip()]

    def _merge_leading_tiny_h2_parts(self, parts: List[str]) -> List[str]:
        """
        After LEGAL_H2_MAX_CHARS splitting, the first paragraph (often only the h2 title line)
        can become its own tiny chunk; merge it into the next part so retrieval does not surface
        a useless heading-only hit.
        """
        if not parts or len(parts) < 2:
            return parts
        min_c = max(100, int(self.min_chunk_chars))
        merged = [p for p in parts if (p or "").strip()]
        safety = 0
        while len(merged) >= 2 and len(merged[0].strip()) < min_c and safety < len(parts) + 5:
            a = merged[0].strip()
            b = merged[1].strip()
            merged = [a + "\n\n" + b] + merged[2:]
            safety += 1
        return merged

    def _expand_oversized_h2_sections(self, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        max_c = self.legal_h2_max_chars
        if max_c <= 0:
            return sections
        expanded: List[Dict[str, Any]] = []
        for sec in sections:
            text = (sec.get("text") or "").strip()
            html = (sec.get("html") or "").strip()
            headers = dict(sec.get("headers") or {})
            if len(text) <= max_c:
                expanded.append(sec)
                continue
            parts = self._split_plain_into_parts(text, max_c)
            parts = self._merge_leading_tiny_h2_parts(parts)
            total = len(parts)
            for pi, part in enumerate(parts):
                h = dict(headers)
                if total > 1:
                    h["h2_part_index"] = str(pi + 1)
                    h["h2_part_total"] = str(total)
                expanded.append(
                    {
                        "html": html,
                        "text": part.strip(),
                        "headers": h,
                    }
                )
        return expanded

    def _strict_split_legal_sections(self, html_doc: str) -> List[Dict[str, Any]]:
        body = self._strip_outer_html(html_doc)
        if not body:
            return []

        headers_to_split: List[Tuple[str, str]] = [
            ("h1", "level_1"),
            ("h2", "level_2"),
        ]
        if self.legal_split_on_h3:
            headers_to_split.append(("h3", "article"))

        splitter = HTMLHeaderTextSplitter(headers_to_split_on=headers_to_split)
        docs = splitter.split_text(body)

        out: List[Dict[str, Any]] = []
        for d in docs:
            md = d.metadata or {}
            section_html = (d.page_content or "").strip()
            if not section_html:
                continue
            level_1 = str(md.get("level_1") or "").strip()
            level_2 = str(md.get("level_2") or "").strip()
            article = str(md.get("article") or "").strip()
            if not self.legal_split_on_h3 and not article:
                article = self._derive_first_h3_plain(section_html)
            header_parts = [x for x in [level_1, level_2, article] if x]
            header_path = " > ".join(header_parts)
            section_text = self._html_to_plain(section_html)
            out.append(
                {
                    "html": section_html,
                    "text": section_text,
                    "headers": {
                        "level_1": level_1,
                        "level_2": level_2,
                        "article": article,
                        "header_path": header_path,
                    },
                }
            )

        return self._expand_oversized_h2_sections(out)

    def _extract_page_from_unit_html(self, unit_html: str) -> Any:
        m = self._RE_PAGE_ATTR.search(unit_html or "")
        if not m:
            return None
        try:
            return int(m.group(1))
        except Exception:
            return None

    # =========================================================
    # Delete old indexed chunks for the same doc_key
    # =========================================================
    def _delete_existing_index(self, doc_key: str) -> None:
        dk = (doc_key or "").strip()
        if not dk:
            return

        rag = self.rag_system
        if rag is None:
            return

        try:
            coll = getattr(rag, "collection", None)
            if coll is not None:
                coll.delete(where={"doc_key": dk})
                print(f"🧹 Deleted previous indexed chunks where doc_key={dk}")
                return
        except Exception as e:
            print("⚠️ Could not delete previous chunks by doc_key:", e)

        try:
            fn = getattr(rag, "delete_doc_key", None)
            if callable(fn):
                fn(dk)
                print(f"🧹 Deleted previous indexed chunks via rag_system.delete_doc_key({dk})")
        except Exception as e:
            print("⚠️ Fallback delete_doc_key failed:", e)

    # =========================================================
    # Main processor
    # =========================================================
    async def process_file(self, file_path: str) -> Dict[str, Any]:
        file_ext = os.path.splitext(file_path)[1].lower()
        request_id = f"{int(time.time() * 1000)}-{os.path.basename(file_path)[:24]}"
        t0 = time.perf_counter()
        self._trace(request_id, "start", t0, f"file_ext={file_ext} path={file_path}")

        try:
            content = ""
            tables_md: List[str] = []
            pdf_pages: List[Dict[str, Any]] = []

            # PDF_HTML_MODE=gemini_file: skip pypdf extraction; HTML comes from Gemini + PDF attachment
            pdf_gemini_file = file_ext == ".pdf" and (
                (os.getenv("PDF_HTML_MODE") or "").strip().lower() == "gemini_file"
            )

            if file_ext in [".png", ".jpg", ".jpeg"]:
                content = self._process_image(file_path)
            elif file_ext == ".pdf":
                if pdf_gemini_file:
                    pdf_pages = []
                    content = ""
                else:
                    pdf_pages = self._process_pdf_pages(file_path)
                    content = "\n\n".join(
                        p["text"] for p in pdf_pages if (p.get("text") or "").strip()
                    ).strip()
            elif file_ext in [".docx", ".doc"]:
                content, tables_md = self._process_docx_with_tables(file_path)
            elif file_ext == ".txt":
                content = self._process_txt(file_path)
            else:
                raise ValueError(f"نوع الملف {file_ext} غير مدعوم")
            self._trace(
                request_id,
                "content_extracted",
                t0,
                f"content_chars={len(content)} tables={len(tables_md)} pdf_pages={len(pdf_pages)}",
            )

            content = (content or "").strip()
            tables_md = [t.strip() for t in (tables_md or []) if (t or "").strip()]

            if not content and not tables_md and not pdf_gemini_file:
                print("⚠️ الملف فارغ بعد المعالجة")
                self._trace(request_id, "stop_empty_content", t0)
                return {"html": "", "converted_txt_path": "", "converted_name": ""}

            stored_filename = os.path.basename(file_path)
            original_name = self._clean_original_name(stored_filename)

            doc_key = (original_name or stored_filename).strip()
            document_id = doc_key

            metadata: Dict[str, Any] = {
                "original_name": original_name,
                "filename": original_name or stored_filename,
                "stored_filename": stored_filename,
                "method": file_ext.replace(".", ""),
                "converted_format": "html",
                "doc_key": doc_key,
            }
            metadata = {k: v for k, v in metadata.items() if v is not None and str(v).strip()}

            combined_parts: List[str] = []
            if content:
                combined_parts.append(content)

            if tables_md:
                table_blocks: List[str] = []
                for idx, md in enumerate(tables_md, start=1):
                    table_blocks.append(f"جدول رقم {idx}:\n{md}")
                combined_parts.append("\n\n".join(table_blocks))

            raw_text_for_llm = "\n\n".join(combined_parts).strip()

            if not self.llm_service:
                raise RuntimeError("llm_service غير متوفر.")
            if pdf_gemini_file:
                if not hasattr(self.llm_service, "to_structured_html_from_pdf"):
                    raise RuntimeError(
                        "llm_service.to_structured_html_from_pdf() غير متوفر. "
                        "تأكد من تعديل llm_service.py وتمريره هنا."
                    )
            elif not hasattr(self.llm_service, "to_structured_html"):
                raise RuntimeError(
                    "llm_service.to_structured_html() غير متوفر. "
                    "تأكد من تعديل llm_service.py وتمريره هنا."
                )

            save_name = original_name or stored_filename

            # =========================================================
            # Build html units
            # For PDF: one unit per page with page_number
            # For others: old segmented behavior with page_number=None
            # =========================================================
            html_units: List[Dict[str, Any]] = []

            if file_ext == ".pdf" and pdf_gemini_file:
                pdf_timeout = max(self.html_llm_timeout, 180)
                print(
                    f"🧩 HTML convert pdf (Gemini file API) timeout={pdf_timeout}s path={file_path}"
                )
                self._trace(
                    request_id,
                    "html_convert_pdf_gemini_file_start",
                    t0,
                    f"timeout={pdf_timeout}s",
                )

                out = ""
                try:
                    out = await asyncio.wait_for(
                        self._maybe_await(
                            self.llm_service.to_structured_html_from_pdf(file_path, save_name)
                        ),
                        timeout=pdf_timeout,
                    )
                except asyncio.TimeoutError:
                    print(f"⏱️ PDF Gemini-file conversion timed out after {pdf_timeout}s")

                inner = self._strip_outer_html(out or "")
                if inner:
                    html_units.append({
                        "page_number": None,
                        "page_text": "",
                        "html": inner,
                    })
                self._trace(
                    request_id,
                    "html_convert_pdf_gemini_file_done",
                    t0,
                    f"html_units={len(html_units)}",
                )

            elif file_ext == ".pdf" and pdf_pages:
                pdf_full_text = "\n\n".join(
                    f"[PAGE {int(p.get('page_number') or 0)}]\n{(p.get('text') or '').strip()}"
                    for p in pdf_pages
                    if (p.get("text") or "").strip()
                ).strip()
                if not pdf_full_text:
                    pdf_full_text = raw_text_for_llm

                if len(pdf_full_text) > self.html_hard_max_chars:
                    pdf_full_text = pdf_full_text[: self.html_hard_max_chars]
                    print(f"✂️ Truncated PDF text to {len(pdf_full_text)} chars (HTML_HARD_MAX_CHARS)")

                pdf_timeout = max(self.html_llm_timeout, 180)
                print(
                    f"🧩 HTML convert pdf (single pass) pages={len(pdf_pages)} chars={len(pdf_full_text)} timeout={pdf_timeout}s"
                )
                self._trace(
                    request_id,
                    "html_convert_pdf_single_start",
                    t0,
                    f"pages={len(pdf_pages)} chars={len(pdf_full_text)} timeout={pdf_timeout}",
                )

                out = ""
                try:
                    out = await asyncio.wait_for(
                        self._maybe_await(self.llm_service.to_structured_html(pdf_full_text)),
                        timeout=pdf_timeout,
                    )
                except asyncio.TimeoutError:
                    print(f"⏱️ PDF single-pass conversion timed out after {pdf_timeout}s")

                inner = self._strip_outer_html(out or "")
                if inner:
                    html_units.append({
                        "page_number": None,
                        "page_text": pdf_full_text,
                        "html": inner,
                    })
                self._trace(request_id, "html_convert_pdf_single_done", t0, f"html_units={len(html_units)}")

            else:
                if len(raw_text_for_llm) > self.html_hard_max_chars:
                    raw_text_for_llm = raw_text_for_llm[: self.html_hard_max_chars]
                    print(f"✂️ Truncated raw text to {len(raw_text_for_llm)} chars (HTML_HARD_MAX_CHARS)")

                segments = self._split_text_segments(raw_text_for_llm)
                print(
                    f"🧩 HTML convert segments={len(segments)} total_chars={len(raw_text_for_llm)} timeout={self.html_llm_timeout}s"
                )
                self._trace(request_id, "html_convert_segments_start", t0, f"segments={len(segments)}")

                for si, seg in enumerate(segments, start=1):
                    seg = (seg or "").strip()
                    if not seg:
                        continue

                    print(f"🚀 Segment {si}/{len(segments)} chars={len(seg)}")

                    try:
                        out = await asyncio.wait_for(
                            self._maybe_await(self.llm_service.to_structured_html(seg)),
                            timeout=self.html_llm_timeout,
                        )
                    except asyncio.TimeoutError:
                        print(f"⏱️ Segment {si} timed out")
                        continue

                    inner = self._strip_outer_html(out or "")
                    if inner:
                        html_units.append({
                            "page_number": None,
                            "page_text": seg,
                            "html": inner,
                        })
                self._trace(request_id, "html_convert_segments_done", t0, f"html_units={len(html_units)}")

            body_inner_all = "\n\n".join(
                (u.get("html") or "").strip()
                for u in html_units
                if (u.get("html") or "").strip()
            ).strip()

            if not body_inner_all:
                print("⚠️ LLM رجّع HTML فارغ")
                self._trace(request_id, "stop_empty_html", t0, "html_units=0")
                return {"html": "", "converted_txt_path": "", "converted_name": ""}

            html = self._wrap_html_document(body_inner_all, title=save_name)
            converted_txt_path, converted_name = self._save_converted_html_as_txt(save_name, html)
            self._trace(request_id, "html_saved", t0, f"converted_name={converted_name}")

            # =========================================================
            # Prepare chunks for indexing (strict legal hierarchy)
            # =========================================================
            chunks_to_index: List[Dict[str, Any]] = []
            table_seq = 0
            for unit in html_units:
                unit_html = (unit.get("html") or "").strip()
                if not unit_html:
                    continue

                page_number = unit.get("page_number")
                if page_number is None:
                    page_number = self._extract_page_from_unit_html(unit_html)
                page_content = (unit.get("page_text") or "").strip()

                normalized_html, table_catalog = self._normalize_html_tables(unit_html)
                sections = self._strict_split_legal_sections(normalized_html)
                if not sections:
                    sections = [
                        {
                            "html": normalized_html,
                            "text": self._html_to_plain(normalized_html),
                            "headers": {
                                "level_1": "",
                                "level_2": "",
                                "article": "",
                                "header_path": "",
                            },
                        }
                    ]

                for sec in sections:
                    sec_text = (sec.get("text") or "").strip()
                    if not sec_text:
                        continue

                    headers = sec.get("headers") or {}
                    section_type = self._section_type_from_headers(headers, sec_text)
                    level_1 = str(headers.get("level_1") or "").strip()
                    level_2 = str(headers.get("level_2") or "").strip()
                    article = str(headers.get("article") or "").strip()

                    # Skip heading-only chunks (e.g. an h2 title indexed alone without body text).
                    # This avoids retrieving tiny title chunks such as "المادة الخامسة والثلاثون".
                    text_norm = re.sub(r"\s+", " ", sec_text).strip()
                    heading_candidates = {
                        re.sub(r"\s+", " ", x).strip()
                        for x in [level_1, level_2, article]
                        if x and re.sub(r"\s+", " ", x).strip()
                    }
                    if heading_candidates and text_norm in heading_candidates:
                        continue

                    is_table = self._is_table_section(sec_text)
                    table_meta = {}
                    if is_table and table_catalog:
                        table_seq += 1
                        table_meta = table_catalog[min(table_seq - 1, len(table_catalog) - 1)]

                    chunks_to_index.append(
                        {
                            "html": (sec.get("html") or "").strip(),
                            "text": sec_text,
                            "page_number": page_number,
                            "page_content": page_content[:4000] if page_content else "",
                            "level_1": level_1,
                            "level_2": level_2,
                            "article": article,
                            "header_path": str(headers.get("header_path") or "").strip(),
                            "h2_part_index": str(headers.get("h2_part_index") or "").strip(),
                            "h2_part_total": str(headers.get("h2_part_total") or "").strip(),
                            "section_type": section_type,
                            "is_table": bool(is_table),
                            "table_id": str(table_meta.get("table_id") or ""),
                            "table_text": str(table_meta.get("table_text") or ""),
                            "table_schema": str(table_meta.get("table_schema") or ""),
                            "parent_header_path": str(headers.get("header_path") or "").strip(),
                        }
                    )

            print(f"🧩 Legal chunks produced: {len(chunks_to_index)}")
            self._trace(request_id, "legal_chunks_built", t0, f"chunks={len(chunks_to_index)}")

            if not chunks_to_index:
                self._trace(request_id, "done_no_chunks", t0)
                return {"html": html, "converted_txt_path": converted_txt_path, "converted_name": converted_name}

            self._delete_existing_index(doc_key)
            self._trace(request_id, "old_index_deleted", t0, f"doc_key={doc_key}")

            if self.rag_system:
                indexed_count = 0
                for idx, item in enumerate(chunks_to_index, start=0):
                    chunk_html = (item.get("html") or "").strip()
                    chunk_plain = (item.get("text") or "").strip()
                    if not chunk_html or not chunk_plain:
                        continue

                    chunk_id = f"{document_id}::chunk_{idx}"

                    chunk_meta = dict(metadata)
                    chunk_meta["converted_file"] = converted_txt_path
                    chunk_meta["skip_chunking"] = True
                    chunk_meta["chunk_index"] = int(idx)
                    chunk_meta["html"] = chunk_html
                    chunk_meta["chunk_type"] = "legal_section"
                    chunk_meta["level_1"] = item.get("level_1") or ""
                    chunk_meta["level_2"] = item.get("level_2") or ""
                    chunk_meta["article"] = item.get("article") or ""
                    chunk_meta["h2"] = item.get("level_2") or ""
                    chunk_meta["h3"] = item.get("article") or ""
                    chunk_meta["header_path"] = item.get("header_path") or ""
                    chunk_meta["parent_header_path"] = item.get("parent_header_path") or ""
                    chunk_meta["section_type"] = item.get("section_type") or "body"
                    chunk_meta["is_table"] = bool(item.get("is_table"))
                    chunk_meta["table_id"] = item.get("table_id") or ""
                    chunk_meta["table_text"] = item.get("table_text") or ""
                    chunk_meta["table_schema"] = item.get("table_schema") or ""
                    if item.get("h2_part_index"):
                        chunk_meta["h2_part_index"] = item.get("h2_part_index") or ""
                    if item.get("h2_part_total"):
                        chunk_meta["h2_part_total"] = item.get("h2_part_total") or ""

                    page_number = item.get("page_number")
                    page_content = (item.get("page_content") or "").strip()

                    if page_number is not None:
                        chunk_meta["page_number"] = int(page_number)

                    if page_content:
                        chunk_meta["page_content"] = page_content[:4000]

                    # Enrich ingested chunk text with document title + h2/h3 context.
                    title = str(
                        chunk_meta.get("original_name")
                        or chunk_meta.get("filename")
                        or chunk_meta.get("name")
                        or ""
                    ).strip()
                    h2_title = str(item.get("level_2") or "").strip()
                    article_h3 = str(item.get("article") or "").strip()
                    chunk_plain_norm = re.sub(r"\s+", " ", chunk_plain).strip()
                    ingest_parts: List[str] = []
                    if title:
                        ingest_parts.append(f"عنوان المستند: {title}")
                    # Keep h2 in indexed text so semantic retrieval can match article/chapter titles directly.
                    if h2_title:
                        h2_norm = re.sub(r"\s+", " ", h2_title).strip()
                        if h2_norm and h2_norm not in chunk_plain_norm:
                            ingest_parts.append(f"العنوان (h2): {h2_title}")
                    if article_h3:
                        h3_norm = re.sub(r"\s+", " ", article_h3).strip()
                        if h3_norm and h3_norm not in chunk_plain_norm:
                            ingest_parts.append(f"العنوان الفرعي (h3): {article_h3}")
                    ingest_parts.append(chunk_plain)
                    ingest_text = "\n".join(ingest_parts).strip()

                    self.rag_system.add_document(
                        content=ingest_text,
                        metadata=chunk_meta,
                        document_id=chunk_id,
                    )
                    indexed_count += 1
                self._trace(request_id, "indexing_done", t0, f"indexed_chunks={indexed_count}")
            else:
                self._trace(request_id, "indexing_skipped_no_rag_system", t0)

            self._trace(request_id, "done", t0)
            return {"html": html, "converted_txt_path": converted_txt_path, "converted_name": converted_name}

        except Exception as e:
            print(f"❌ DocumentProcessor error: {e}")
            self._trace(request_id, "error", t0, f"error={e}")
            raise

    # =========================================================
    # OCR deps (optional)
    # =========================================================
    def _try_import_ocr(self):
        if not self.enable_ocr:
            return None

        try:
            import pytesseract
            from PIL import Image
            from pdf2image import convert_from_path

            if self.tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd

            return pytesseract, Image, convert_from_path
        except Exception as e:
            print("⚠️ OCR enabled لكن مكتبات OCR غير جاهزة:", e)
            return None

    # =========================================================
    # File handlers
    # =========================================================
    def _process_image(self, file_path: str) -> str:
        if not self.enable_ocr:
            return ""

        ocr = self._try_import_ocr()
        if not ocr:
            return ""

        pytesseract, Image, _convert_from_path = ocr
        text = pytesseract.image_to_string(Image.open(file_path), lang="ara+eng")
        return self._clean_text_keep_tables(text)

    def _process_pdf_pages(self, file_path: str) -> List[Dict[str, Any]]:
        reader = PdfReader(file_path)
        raw_pages: List[Dict[str, Any]] = []

        for i, page in enumerate(reader.pages, start=1):
            extracted = page.extract_text() or ""
            cleaned = self._clean_text_keep_tables(extracted)

            if cleaned.strip():
                raw_pages.append({
                    "page_number": i,
                    "text": cleaned,
                })

        # OCR fallback only if no extracted pages found
        if self.enable_ocr and not raw_pages:
            ocr = self._try_import_ocr()
            if ocr:
                pytesseract, _Image, convert_from_path = ocr
                try:
                    images = convert_from_path(file_path)
                    for i, img in enumerate(images, start=1):
                        ocr_text = pytesseract.image_to_string(img, lang="ara+eng")
                        cleaned = self._clean_text_keep_tables(ocr_text)
                        if cleaned.strip():
                            raw_pages.append({
                                "page_number": i,
                                "text": cleaned,
                            })
                except Exception as e:
                    print("⚠️ PDF OCR failed:", e)

        if not raw_pages:
            return []

        repeated_top, repeated_bottom = self._collect_repeated_edge_lines(
            raw_pages,
            edge_line_count=3,
            min_repeat=2,
        )

        pages: List[Dict[str, Any]] = []
        for page in raw_pages:
            cleaned_page = self._remove_header_footer_from_page_text(
                page.get("text") or "",
                repeated_top=repeated_top,
                repeated_bottom=repeated_bottom,
                edge_line_count=3,
            )
            if cleaned_page.strip():
                pages.append({
                    "page_number": page.get("page_number"),
                    "text": cleaned_page,
                })

        return pages

    def _process_docx_with_tables(self, file_path: str) -> Tuple[str, List[str]]:
        doc = DocxDocument(file_path)

        paras: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                paras.append(t)

        tables_md: List[str] = []
        for table in doc.tables:
            md = self._docx_table_to_markdown(table).strip()
            if md:
                tables_md.append(md)

        content = "\n".join(paras).strip()
        return self._clean_text_keep_tables(content), tables_md

    def _process_txt(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return self._clean_text_keep_tables(f.read())

    # =========================================================
    # DOCX table -> Markdown
    # =========================================================
    def _docx_table_to_markdown(self, table) -> str:
        rows: List[List[str]] = []
        max_cols = 0

        for row in table.rows:
            cells = []
            for cell in row.cells:
                txt = (cell.text or "").strip()
                txt = re.sub(r"\s+", " ", txt).strip()
                cells.append(txt)
            max_cols = max(max_cols, len(cells))
            rows.append(cells)

        if not rows or max_cols == 0:
            return ""

        norm_rows: List[List[str]] = []
        for r in rows:
            rr = r + [""] * (max_cols - len(r))
            norm_rows.append(rr)

        def esc(x: str) -> str:
            x = (x or "").strip()
            return x.replace("|", "\\|")

        def _norm_cell(s: str) -> str:
            s = (s or "").strip()
            s = re.sub(r"\s+", " ", s)
            return s

        def _looks_like_caption_row(row_vals: List[str]) -> bool:
            cleaned = [_norm_cell(c) for c in row_vals if _norm_cell(c)]
            if not cleaned:
                return False
            if len(set(cleaned)) == 1:
                return True
            if len(cleaned) == 1 and len(cleaned[0]) >= 10:
                return True
            return False

        caption = ""
        data_rows = norm_rows

        if _looks_like_caption_row(norm_rows[0]):
            caption = _norm_cell(norm_rows[0][0]) or next((c for c in norm_rows[0] if _norm_cell(c)), "")
            data_rows = norm_rows[1:] if len(norm_rows) > 1 else []

        if not data_rows:
            return f"**{caption}**" if caption else ""

        if max_cols == 2:
            header = ["المعدل", "التقدير"]
            body = data_rows
        else:
            header = data_rows[0]
            body = data_rows[1:] if len(data_rows) > 1 else []

        header_line = "| " + " | ".join(esc(x) for x in header) + " |"
        sep_line = "| " + " | ".join(["---"] * max_cols) + " |"

        body_lines = []
        for r in body:
            body_lines.append("| " + " | ".join(esc(x) for x in r) + " |")

        if caption:
            return ("**" + esc(caption) + "**\n\n" + "\n".join([header_line, sep_line] + body_lines)).strip()

        return "\n".join([header_line, sep_line] + body_lines).strip()

    # =========================================================
    # Cleaning
    # =========================================================
    def _clean_text_keep_tables(self, text: str) -> str:
        if not text:
            return ""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()